"""
R&D Report Document Generator
Generates HMRC-compliant Word documents from R&D activities
"""

from typing import List, Dict
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from rd_classifier import RDActivity
import os


class RDReportGenerator:
    """
    Generates professional R&D tax relief documentation in Word format
    """
    
    def __init__(self, company_name: str = ""):
        """
        Initialize report generator
        
        Args:
            company_name: Name of company for report header
        """
        self.company_name = company_name
    
    def generate_report(
        self,
        activities: List[RDActivity],
        repo_name: str,
        period_start: datetime,
        period_end: datetime,
        output_path: str = "rd_technical_report.docx"
    ) -> str:
        """
        Generate complete R&D tax relief technical report
        
        Args:
            activities: List of R&D activities to document
            repo_name: Repository analyzed
            period_start: Start date of analysis period
            period_end: End date of analysis period
            output_path: Where to save the document
            
        Returns:
            Path to generated document
        """
        # Create document
        doc = Document()
        
        # Set up document styling
        self._apply_document_styles(doc)
        
        # Add cover page
        self._add_cover_page(doc, repo_name, period_start, period_end)
        
        # Add executive summary
        self._add_executive_summary(doc, activities, period_start, period_end)
        
        # Add table of contents placeholder
        self._add_toc_placeholder(doc)
        
        # Add methodology section
        self._add_methodology(doc)
        
        # Add each R&D activity
        self._add_rd_activities(doc, activities)
        
        # Add appendix with evidence
        self._add_evidence_appendix(doc, activities)
        
        # Add conclusion
        self._add_conclusion(doc, activities)
        
        # Save document
        doc.save(output_path)
        print(f"✓ Report generated: {output_path}")
        
        return output_path
    
    def _apply_document_styles(self, doc: Document) -> None:
        """Apply consistent styling to document"""
        
        # Define styles for headings, body text, etc.
        styles = doc.styles
        
        # Heading 1 style
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 128)
        
        # Normal text
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
    
    def _add_cover_page(
        self,
        doc: Document,
        repo_name: str,
        period_start: datetime,
        period_end: datetime
    ) -> None:
        """Add cover page"""
        
        # Title
        title = doc.add_heading('R&D Tax Relief', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_heading('Technical Report', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Company info
        if self.company_name:
            p = doc.add_paragraph()
            p.add_run(f'Company: {self.company_name}').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Project info
        p = doc.add_paragraph()
        p.add_run(f'Project: {repo_name}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Period
        p = doc.add_paragraph()
        period_text = f'Period: {period_start.strftime("%B %d, %Y")} - {period_end.strftime("%B %d, %Y")}'
        p.add_run(period_text).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generated date
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
    
    def _add_executive_summary(
        self,
        doc: Document,
        activities: List[RDActivity],
        period_start: datetime,
        period_end: datetime
    ) -> None:
        """Add executive summary"""
        
        doc.add_heading('Executive Summary', level=1)
        
        # Summary stats
        total_activities = len(activities)
        high_confidence = len([a for a in activities if a.confidence_score >= 75])
        avg_confidence = sum(a.confidence_score for a in activities) / len(activities) if activities else 0
        
        summary_text = f"""
        This technical report documents {total_activities} qualifying R&D activities undertaken 
        between {period_start.strftime("%B %Y")} and {period_end.strftime("%B %Y")}. 
        
        These activities meet the criteria for Corporation Tax R&D Relief as defined by HMRC 
        guidance, specifically addressing technological uncertainties through systematic 
        investigation to achieve advances in the field of science and technology.
        
        Of the {total_activities} activities identified, {high_confidence} demonstrate high 
        confidence (≥75%) in meeting HMRC criteria, with an average confidence score of 
        {avg_confidence:.1f}%.
        
        All activities documented below include:
        • Clear identification of technological uncertainty
        • Evidence of systematic investigation
        • Demonstration of technical advance achieved
        • Supporting evidence from version control history
        """
        
        doc.add_paragraph(summary_text.strip())
        doc.add_page_break()
    
    def _add_toc_placeholder(self, doc: Document) -> None:
        """Add table of contents placeholder"""
        
        doc.add_heading('Table of Contents', level=1)
        p = doc.add_paragraph()
        p.add_run('[Table of Contents - Generate in Word via References > Table of Contents]').italic = True
        doc.add_page_break()
    
    def _add_methodology(self, doc: Document) -> None:
        """Add methodology section"""
        
        doc.add_heading('Methodology', level=1)
        
        methodology_text = """
        This analysis was conducted using an automated R&D identification system that:
        
        1. Extracted development activity from version control (GitHub) including commits, 
           pull requests, and documentation
           
        2. Analyzed code changes and technical descriptions against HMRC R&D criteria using 
           advanced AI classification
           
        3. Retrieved relevant HMRC guidance using semantic search to ensure accurate 
           criterion matching
           
        4. Assessed each activity for:
           • Technological uncertainty (problems without readily available solutions)
           • Systematic investigation (structured problem-solving approach)
           • Technical advance (new capabilities or knowledge)
           
        5. Generated technical narratives suitable for HMRC submission
        
        Each activity includes a confidence score indicating the strength of evidence for 
        R&D qualification. Higher scores indicate stronger alignment with HMRC criteria.
        """
        
        doc.add_paragraph(methodology_text.strip())
        doc.add_page_break()
    
    def _add_rd_activities(self, doc: Document, activities: List[RDActivity]) -> None:
        """Add detailed R&D activities section"""
        
        doc.add_heading('R&D Activities', level=1)
        
        for i, activity in enumerate(activities, 1):
            # Activity heading
            doc.add_heading(f'{i}. {activity.title}', level=2)
            
            # Confidence indicator
            confidence_para = doc.add_paragraph()
            confidence_run = confidence_para.add_run(
                f'Confidence Score: {activity.confidence_score:.0f}%'
            )
            if activity.confidence_score >= 75:
                confidence_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif activity.confidence_score >= 50:
                confidence_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                confidence_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Timeframe
            doc.add_paragraph(f'Timeframe: {activity.timeframe}')
            
            # Description
            doc.add_heading('Description', level=3)
            doc.add_paragraph(activity.description)
            
            # Technological Uncertainty (HMRC Criterion 1)
            doc.add_heading('Technological Uncertainty', level=3)
            doc.add_paragraph(activity.technological_uncertainty)
            
            # Systematic Investigation (HMRC Criterion 2)
            doc.add_heading('Systematic Investigation', level=3)
            doc.add_paragraph(activity.systematic_investigation)
            
            # Technical Advance (HMRC Criterion 3)
            doc.add_heading('Technical Advance', level=3)
            doc.add_paragraph(activity.technical_advance)
            
            # Supporting Evidence
            doc.add_heading('Supporting Evidence', level=3)
            if activity.commits:
                doc.add_paragraph(f'Commits: {len(activity.commits)} technical commits')
                # List first few commits
                for commit_sha in activity.commits[:3]:
                    doc.add_paragraph(f'• {commit_sha[:8]}', style='List Bullet')
            
            if activity.pull_requests:
                doc.add_paragraph(f'Pull Requests: {len(activity.pull_requests)} PRs')
                for pr_num in activity.pull_requests[:3]:
                    doc.add_paragraph(f'• PR #{pr_num}', style='List Bullet')
            
            # Separator
            doc.add_paragraph()
            if i < len(activities):
                doc.add_paragraph('─' * 80)
                doc.add_paragraph()
    
    def _add_evidence_appendix(self, doc: Document, activities: List[RDActivity]) -> None:
        """Add appendix with detailed evidence"""
        
        doc.add_page_break()
        doc.add_heading('Appendix: Supporting Evidence', level=1)
        
        for i, activity in enumerate(activities, 1):
            doc.add_heading(f'Activity {i}: {activity.title}', level=2)
            
            # Commit evidence
            if activity.commits:
                doc.add_heading('Commits', level=3)
                for commit_sha in activity.commits:
                    doc.add_paragraph(f'• {commit_sha}', style='List Bullet')
            
            # PR evidence
            if activity.pull_requests:
                doc.add_heading('Pull Requests', level=3)
                for pr_num in activity.pull_requests:
                    doc.add_paragraph(f'• PR #{pr_num}', style='List Bullet')
            
            doc.add_paragraph()
    
    def _add_conclusion(self, doc: Document, activities: List[RDActivity]) -> None:
        """Add conclusion section"""
        
        doc.add_page_break()
        doc.add_heading('Conclusion', level=1)
        
        conclusion_text = f"""
        This technical report documents {len(activities)} qualifying R&D activities that meet 
        HMRC's criteria for Corporation Tax R&D Relief. Each activity demonstrates:
        
        • Technological uncertainty that required resolution
        • Systematic and scientific approach to investigation
        • Technical advances in the relevant field
        • Strong supporting evidence from development records
        
        All activities are supported by version control evidence, technical documentation, 
        and detailed narratives that comply with HMRC reporting requirements.
        
        This documentation provides a robust foundation for an R&D tax relief claim and 
        includes sufficient detail to support HMRC compliance and potential audit requirements.
        """
        
        doc.add_paragraph(conclusion_text.strip())


# Example usage
if __name__ == "__main__":
    from datetime import datetime, timedelta
    
    # Sample R&D activities (in real use, these come from the classifier)
    sample_activities = [
        RDActivity(
            id="activity_1",
            title="Novel Machine Learning Algorithm for Query Optimization",
            description="""
            Developed a new reinforcement learning-based approach to database query optimization 
            that significantly outperformed existing rule-based and cost-based optimizers. The 
            challenge was handling non-stationary workload patterns in high-throughput environments.
            """,
            timeframe="Q3 2024",
            technological_uncertainty="""
            At the project's outset, there was significant technological uncertainty regarding 
            whether reinforcement learning could effectively adapt to rapidly changing query 
            patterns in production database systems. Existing literature provided no clear 
            solution for this specific use case, and competent professionals in the field could 
            not readily deduce an optimal approach.
            """,
            systematic_investigation="""
            We undertook a systematic investigation including: (1) Literature review of RL 
            approaches in database optimization, (2) Design and implementation of multiple RL 
            agent architectures, (3) Controlled experimentation with various reward functions, 
            (4) Iterative refinement based on measured performance metrics, (5) A/B testing 
            against baseline optimizers.
            """,
            technical_advance="""
            This work achieved a technical advance by demonstrating that RL agents with 
            adaptive reward shaping can achieve 40% better query throughput than traditional 
            optimizers in non-stationary workloads. This represents new knowledge in the field 
            of database optimization and has been submitted for publication.
            """,
            commits=["a1b2c3d", "e4f5g6h", "i7j8k9l"],
            pull_requests=[123, 145, 167],
            confidence_score=85.0,
            created_at="2024-09-15"
        )
    ]
    
    # Generate report
    generator = RDReportGenerator(company_name="Example Tech Ltd")
    
    report_path = generator.generate_report(
        activities=sample_activities,
        repo_name="example/ml-optimizer",
        period_start=datetime.now() - timedelta(days=90),
        period_end=datetime.now(),
        output_path="sample_rd_report.docx"
    )
    
    print(f"Sample report generated: {report_path}")
